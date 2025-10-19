"""
DOCX Text Extractor

Extracts text from Microsoft Word (.docx) files using python-docx library.
Handles paragraphs, headings, tables, and preserves document structure.
"""

import io
import logging
from typing import Tuple, Optional

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    raise ImportError(
        "python-docx is required for DOCX extraction. "
        "Install with: pip install python-docx"
    )

from ..document_processor import TextExtractor

logger = logging.getLogger(__name__)


class DOCXExtractor(TextExtractor):
    """
    Extracts text content from DOCX files

    Features:
    - Paragraph and heading extraction
    - Table content extraction
    - Structure preservation with separators
    - Error handling for corrupted files
    """

    def __init__(
        self,
        paragraph_separator: str = "\n",
        section_separator: str = "\n\n",
        table_cell_separator: str = " | ",
        table_row_separator: str = "\n"
    ):
        """
        Initialize DOCX extractor

        Args:
            paragraph_separator: String between paragraphs (default: newline)
            section_separator: String between major sections (default: double newline)
            table_cell_separator: String between table cells (default: " | ")
            table_row_separator: String between table rows (default: newline)
        """
        self.paragraph_separator = paragraph_separator
        self.section_separator = section_separator
        self.table_cell_separator = table_cell_separator
        self.table_row_separator = table_row_separator

    def extract(self, content: bytes, filename: str) -> str:
        """
        Extract text from DOCX file

        Args:
            content: Raw DOCX file bytes
            filename: Original filename (for error messages)

        Returns:
            Extracted text with structure preserved

        Raises:
            ValueError: If DOCX is corrupted or cannot be read
        """
        try:
            # Create Document from bytes
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            logger.debug(f"Processing DOCX file: {filename}")

            text_parts = []

            # Extract paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Skip empty paragraphs
                    text_parts.append(text)
                    paragraph_count += 1

            logger.debug(f"Extracted {paragraph_count} paragraphs from {filename}")

            # Extract tables
            table_count = 0
            if doc.tables:
                for table in doc.tables:
                    table_text = self._extract_table(table)
                    if table_text:
                        # Add section separator before table
                        text_parts.append(self.section_separator)
                        text_parts.append(table_text)
                        table_count += 1

            logger.debug(f"Extracted {table_count} tables from {filename}")

            if not text_parts:
                raise ValueError(
                    f"No text could be extracted from DOCX '{filename}'. "
                    "File may be empty or corrupted."
                )

            # Combine all parts
            full_text = self.paragraph_separator.join(text_parts)

            logger.info(
                f"Successfully extracted {len(full_text)} characters "
                f"from {filename} ({paragraph_count} paragraphs, {table_count} tables)"
            )

            return full_text

        except PackageNotFoundError as e:
            raise ValueError(
                f"Failed to read DOCX '{filename}': {str(e)}. File may be corrupted or not a valid DOCX."
            )
        except Exception as e:
            if isinstance(e, ValueError):
                raise  # Re-raise our custom errors
            logger.error(f"Unexpected error extracting text from {filename}: {str(e)}")
            raise ValueError(
                f"Failed to extract text from DOCX '{filename}': {str(e)}"
            )

    def _extract_table(self, table) -> str:
        """
        Extract text from a table

        Args:
            table: python-docx Table object

        Returns:
            Formatted table text
        """
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:  # Only add non-empty cells
                    cells.append(cell_text)

            if cells:  # Only add rows with content
                rows.append(self.table_cell_separator.join(cells))

        return self.table_row_separator.join(rows)

    def validate(self, content: bytes) -> Tuple[bool, Optional[str]]:
        """
        Validate that content is a valid DOCX file

        Args:
            content: Raw file bytes

        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check file signature (DOCX is a ZIP file with specific structure)
            # DOCX files start with PK (ZIP signature)
            if not content.startswith(b'PK'):
                return False, "File does not appear to be a valid DOCX (missing ZIP header)"

            # Try to open with python-docx
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            # Basic validation - try to access document properties
            _ = doc.paragraphs
            _ = doc.tables

            return True, None

        except PackageNotFoundError as e:
            return False, f"Invalid or corrupted DOCX: {str(e)}"
        except Exception as e:
            return False, f"DOCX validation failed: {str(e)}"

    def get_metadata(self, content: bytes) -> dict:
        """
        Extract DOCX metadata (optional helper method)

        Args:
            content: Raw DOCX file bytes

        Returns:
            Dictionary with metadata (title, author, created, etc.)
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            metadata = {}

            # Get core properties
            core_props = doc.core_properties
            if core_props:
                metadata['title'] = core_props.title or ''
                metadata['author'] = core_props.author or ''
                metadata['subject'] = core_props.subject or ''
                metadata['keywords'] = core_props.keywords or ''
                metadata['comments'] = core_props.comments or ''
                metadata['created'] = str(core_props.created) if core_props.created else ''
                metadata['modified'] = str(core_props.modified) if core_props.modified else ''
                metadata['last_modified_by'] = core_props.last_modified_by or ''

            # Add document statistics
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            metadata['section_count'] = len(doc.sections)

            return metadata

        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {str(e)}")
            return {}
