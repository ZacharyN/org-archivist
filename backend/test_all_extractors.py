"""
Test script for all document extractors (PDF, DOCX, TXT)
Creates sample files and tests extraction functionality
"""

import io
import sys
from pathlib import Path

# Add backend directory to path
sys.path.insert(0, str(Path(__file__).parent))

# Create test files
def create_test_files():
    """Create sample PDF, DOCX, and TXT files for testing"""

    # 1. Create test PDF
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("Test PDF Document", styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph("This is a test paragraph in the PDF.", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph("This is another paragraph on page 1.", styles['Normal']))
        story.append(Spacer(1, 200))
        story.append(Paragraph("This paragraph is on page 2.", styles['Normal']))

        doc.build(story)
        pdf_content = pdf_buffer.getvalue()
        print("[OK] Created test PDF")
    except ImportError:
        print("[SKIP] PDF creation (reportlab not installed)")
        pdf_content = None

    # 2. Create test DOCX
    try:
        from docx import Document
        from docx.shared import Pt

        docx_buffer = io.BytesIO()
        doc = Document()

        # Add heading
        doc.add_heading('Test DOCX Document', level=1)

        # Add paragraphs
        doc.add_paragraph('This is a test paragraph in the DOCX file.')
        doc.add_paragraph('Here is another paragraph with some content.')

        # Add a table
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light Grid Accent 1'

        # Populate table
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Header 1'
        header_cells[1].text = 'Header 2'
        header_cells[2].text = 'Header 3'

        row1_cells = table.rows[1].cells
        row1_cells[0].text = 'Row 1, Col 1'
        row1_cells[1].text = 'Row 1, Col 2'
        row1_cells[2].text = 'Row 1, Col 3'

        row2_cells = table.rows[2].cells
        row2_cells[0].text = 'Row 2, Col 1'
        row2_cells[1].text = 'Row 2, Col 2'
        row2_cells[2].text = 'Row 2, Col 3'

        # Add another paragraph
        doc.add_paragraph('This is a paragraph after the table.')

        doc.save(docx_buffer)
        docx_content = docx_buffer.getvalue()
        print("[OK] Created test DOCX")
    except ImportError:
        print("[SKIP] DOCX creation (python-docx not installed)")
        docx_content = None

    # 3. Create test TXT files with different encodings
    txt_utf8_content = """Test Text Document (UTF-8)

This is a plain text file.
It has multiple lines.

Here's another paragraph with special characters: café, naïve, résumé

End of document.""".encode('utf-8')

    txt_latin1_content = "Test Latin-1 Document\n\nCafé, naïve, résumé".encode('latin-1')

    print("[OK] Created test TXT files")

    return pdf_content, docx_content, txt_utf8_content, txt_latin1_content


def test_extractors():
    """Test all three extractors"""

    from app.services.extractors.pdf_extractor import PDFExtractor
    from app.services.extractors.docx_extractor import DOCXExtractor
    from app.services.extractors.txt_extractor import TXTExtractor

    print("\n" + "="*60)
    print("Testing Document Extractors")
    print("="*60 + "\n")

    # Create test files
    pdf_content, docx_content, txt_utf8, txt_latin1 = create_test_files()

    # Test PDF Extractor
    print("\n--- Testing PDF Extractor ---")
    if pdf_content:
        try:
            pdf_extractor = PDFExtractor()

            # Validate
            is_valid, error = pdf_extractor.validate(pdf_content)
            print(f"Validation: {'[PASS]' if is_valid else '[FAIL]'}")
            if error:
                print(f"  Error: {error}")

            # Extract
            text = pdf_extractor.extract(pdf_content, "test.pdf")
            print(f"Extracted {len(text)} characters")
            print(f"Preview: {text[:150]}...")

            # Metadata
            metadata = pdf_extractor.get_metadata(pdf_content)
            print(f"Metadata: {metadata}")

            print("[OK] PDF Extractor: SUCCESS")
        except Exception as e:
            print(f"[FAIL] PDF Extractor: FAILED - {str(e)}")
    else:
        print("[SKIP] PDF Extractor: SKIPPED")

    # Test DOCX Extractor
    print("\n--- Testing DOCX Extractor ---")
    if docx_content:
        try:
            docx_extractor = DOCXExtractor()

            # Validate
            is_valid, error = docx_extractor.validate(docx_content)
            print(f"Validation: {'[PASS]' if is_valid else '[FAIL]'}")
            if error:
                print(f"  Error: {error}")

            # Extract
            text = docx_extractor.extract(docx_content, "test.docx")
            print(f"Extracted {len(text)} characters")
            print(f"Preview: {text[:200]}...")

            # Check if table was extracted
            if "Header 1" in text and "Row 1, Col 1" in text:
                print("[OK] Table extraction: SUCCESS")
            else:
                print("[FAIL] Table extraction: FAILED")

            # Metadata
            metadata = docx_extractor.get_metadata(docx_content)
            print(f"Metadata: {metadata}")

            print("[OK] DOCX Extractor: SUCCESS")
        except Exception as e:
            print(f"[FAIL] DOCX Extractor: FAILED - {str(e)}")
    else:
        print("[SKIP] DOCX Extractor: SKIPPED")

    # Test TXT Extractor (UTF-8)
    print("\n--- Testing TXT Extractor (UTF-8) ---")
    try:
        txt_extractor = TXTExtractor()

        # Validate
        is_valid, error = txt_extractor.validate(txt_utf8)
        print(f"Validation: {'[PASS]' if is_valid else '[FAIL]'}")
        if error:
            print(f"  Error: {error}")

        # Extract
        text = txt_extractor.extract(txt_utf8, "test_utf8.txt")
        print(f"Extracted {len(text)} characters")
        print(f"Preview: {text[:150]}...")

        # Check special characters
        if "café" in text and "naïve" in text and "résumé" in text:
            print("[OK] Special character handling: SUCCESS")
        else:
            print("[FAIL] Special character handling: FAILED")

        # Encoding info
        encoding_info = txt_extractor.get_encoding_info(txt_utf8)
        print(f"Encoding info: {encoding_info}")

        print("[OK] TXT Extractor (UTF-8): SUCCESS")
    except Exception as e:
        print(f"[FAIL] TXT Extractor (UTF-8): FAILED - {str(e)}")

    # Test TXT Extractor (Latin-1)
    print("\n--- Testing TXT Extractor (Latin-1) ---")
    try:
        txt_extractor = TXTExtractor()

        # Validate
        is_valid, error = txt_extractor.validate(txt_latin1)
        print(f"Validation: {'[PASS]' if is_valid else '[FAIL]'}")
        if error:
            print(f"  Error: {error}")

        # Extract
        text = txt_extractor.extract(txt_latin1, "test_latin1.txt")
        print(f"Extracted {len(text)} characters")
        print(f"Preview: {text}")

        # Encoding info
        encoding_info = txt_extractor.get_encoding_info(txt_latin1)
        print(f"Encoding info: {encoding_info}")

        print("[OK] TXT Extractor (Latin-1): SUCCESS")
    except Exception as e:
        print(f"[FAIL] TXT Extractor (Latin-1): FAILED - {str(e)}")

    # Test error handling
    print("\n--- Testing Error Handling ---")

    # Invalid PDF
    try:
        pdf_extractor = PDFExtractor()
        pdf_extractor.extract(b"not a pdf", "invalid.pdf")
        print("[FAIL] PDF error handling: FAILED (should have raised error)")
    except ValueError as e:
        print(f"[OK] PDF error handling: SUCCESS - {str(e)[:50]}...")

    # Invalid DOCX
    try:
        docx_extractor = DOCXExtractor()
        docx_extractor.extract(b"not a docx", "invalid.docx")
        print("[FAIL] DOCX error handling: FAILED (should have raised error)")
    except ValueError as e:
        print(f"[OK] DOCX error handling: SUCCESS - {str(e)[:50]}...")

    # Empty TXT
    try:
        txt_extractor = TXTExtractor()
        txt_extractor.extract(b"", "empty.txt")
        print("[FAIL] TXT error handling: FAILED (should have raised error)")
    except ValueError as e:
        print(f"[OK] TXT error handling: SUCCESS - {str(e)[:50]}...")

    print("\n" + "="*60)
    print("All Tests Complete!")
    print("="*60)


if __name__ == "__main__":
    test_extractors()
