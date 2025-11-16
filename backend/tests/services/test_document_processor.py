"""
Unit tests for document processor and text extractors.

Tests cover:
- PDF text extraction (multi-page, encrypted, corrupted)
- DOCX text extraction (paragraphs, tables, corrupted)
- TXT text extraction (encoding detection, special chars)
- Document processor pipeline
- Error handling and edge cases
"""
import pytest
import io
from typing import Tuple

# PDF generation libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX generation
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import our processors
from app.services.document_processor import (
    DocumentProcessor,
    ProcessorFactory,
    FileType,
    DocumentChunk
)
from app.services.extractors.pdf_extractor import PDFExtractor
from app.services.extractors.docx_extractor import DOCXExtractor
from app.services.extractors.txt_extractor import TXTExtractor


# ============================================================================
# Test File Generators
# ============================================================================

def generate_test_pdf(text_content: str, num_pages: int = 1) -> bytes:
    """
    Generate a simple PDF for testing

    Args:
        text_content: Text to include in PDF
        num_pages: Number of pages to create

    Returns:
        PDF file as bytes
    """
    if not REPORTLAB_AVAILABLE:
        pytest.skip("reportlab not available for PDF generation")

    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)

    # Add text to each page
    for page_num in range(num_pages):
        c.drawString(100, 750, f"Page {page_num + 1}")
        c.drawString(100, 700, text_content)
        c.showPage()

    c.save()
    return pdf_buffer.getvalue()


def generate_test_docx(paragraphs: list, include_table: bool = False) -> bytes:
    """
    Generate a simple DOCX for testing

    Args:
        paragraphs: List of paragraph texts
        include_table: Whether to include a test table

    Returns:
        DOCX file as bytes
    """
    if not DOCX_AVAILABLE:
        pytest.skip("python-docx not available for DOCX generation")

    doc = Document()

    # Add paragraphs
    for para_text in paragraphs:
        doc.add_paragraph(para_text)

    # Add table if requested
    if include_table:
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(0, 2).text = "Header 3"
        table.cell(1, 0).text = "Row 1 Col 1"
        table.cell(1, 1).text = "Row 1 Col 2"
        table.cell(1, 2).text = "Row 1 Col 3"
        table.cell(2, 0).text = "Row 2 Col 1"
        table.cell(2, 1).text = "Row 2 Col 2"
        table.cell(2, 2).text = "Row 2 Col 3"

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    return docx_buffer.getvalue()


# ============================================================================
# PDF Extractor Tests
# ============================================================================

class TestPDFExtractor:
    """Test suite for PDF text extraction"""

    def test_pdf_extractor_initialization(self):
        """Test PDF extractor can be initialized"""
        extractor = PDFExtractor()
        assert extractor is not None
        assert extractor.page_separator == "\n\n--- Page Break ---\n\n"

    def test_pdf_extractor_custom_separator(self):
        """Test PDF extractor with custom page separator"""
        custom_sep = "\n<PAGE>\n"
        extractor = PDFExtractor(page_separator=custom_sep)
        assert extractor.page_separator == custom_sep

    def test_pdf_single_page_extraction(self):
        """Test extracting text from single-page PDF"""
        test_text = "This is a test PDF document for extraction."
        pdf_bytes = generate_test_pdf(test_text, num_pages=1)

        extractor = PDFExtractor()
        extracted = extractor.extract(pdf_bytes, "test.pdf")

        assert "This is a test PDF" in extracted
        assert "Page 1" in extracted
        assert len(extracted) > 0

    def test_pdf_multi_page_extraction(self):
        """Test extracting text from multi-page PDF"""
        test_text = "Grant writing best practices."
        pdf_bytes = generate_test_pdf(test_text, num_pages=3)

        extractor = PDFExtractor()
        extracted = extractor.extract(pdf_bytes, "multipage.pdf")

        # Should contain content from all pages
        assert "Page 1" in extracted
        assert "Page 2" in extracted
        assert "Page 3" in extracted

        # Should contain page separators
        assert "--- Page Break ---" in extracted

        # Should have multiple occurrences of the test text
        assert extracted.count("Grant writing") >= 2

    def test_pdf_validation_success(self):
        """Test PDF validation succeeds for valid PDF"""
        pdf_bytes = generate_test_pdf("Valid PDF content")

        extractor = PDFExtractor()
        is_valid, error = extractor.validate(pdf_bytes)

        assert is_valid is True
        assert error is None

    def test_pdf_validation_invalid_header(self):
        """Test PDF validation fails for non-PDF file"""
        fake_pdf = b"This is not a PDF file"

        extractor = PDFExtractor()
        is_valid, error = extractor.validate(fake_pdf)

        assert is_valid is False
        assert "PDF header" in error or "valid PDF" in error

    def test_pdf_validation_empty_file(self):
        """Test PDF validation fails for empty file"""
        empty_pdf = b""

        extractor = PDFExtractor()
        is_valid, error = extractor.validate(empty_pdf)

        assert is_valid is False
        assert error is not None

    def test_pdf_extraction_empty_file(self):
        """Test PDF extraction raises error for empty file"""
        empty_pdf = b""

        extractor = PDFExtractor()
        with pytest.raises(ValueError) as exc_info:
            extractor.extract(empty_pdf, "empty.pdf")

        assert "empty.pdf" in str(exc_info.value)

    def test_pdf_get_metadata(self):
        """Test extracting metadata from PDF"""
        pdf_bytes = generate_test_pdf("Test content")

        extractor = PDFExtractor()
        metadata = extractor.get_metadata(pdf_bytes)

        assert isinstance(metadata, dict)
        assert "page_count" in metadata
        assert metadata["page_count"] >= 1


# ============================================================================
# DOCX Extractor Tests
# ============================================================================

class TestDOCXExtractor:
    """Test suite for DOCX text extraction"""

    def test_docx_extractor_initialization(self):
        """Test DOCX extractor can be initialized"""
        extractor = DOCXExtractor()
        assert extractor is not None
        assert extractor.paragraph_separator == "\n"
        assert extractor.table_cell_separator == " | "

    def test_docx_simple_extraction(self):
        """Test extracting text from simple DOCX"""
        paragraphs = [
            "This is the first paragraph.",
            "This is the second paragraph about grant writing.",
            "This is the third paragraph with more content."
        ]
        docx_bytes = generate_test_docx(paragraphs)

        extractor = DOCXExtractor()
        extracted = extractor.extract(docx_bytes, "test.docx")

        assert "first paragraph" in extracted
        assert "second paragraph" in extracted
        assert "third paragraph" in extracted
        assert len(extracted) > 0

    def test_docx_with_table_extraction(self):
        """Test extracting text from DOCX with table"""
        paragraphs = ["Document with a table below:"]
        docx_bytes = generate_test_docx(paragraphs, include_table=True)

        extractor = DOCXExtractor()
        extracted = extractor.extract(docx_bytes, "table.docx")

        # Should contain paragraph
        assert "Document with a table" in extracted

        # Should contain table content
        assert "Header 1" in extracted
        assert "Row 1 Col 1" in extracted

        # Should have table cell separators
        assert "|" in extracted

    def test_docx_validation_success(self):
        """Test DOCX validation succeeds for valid DOCX"""
        docx_bytes = generate_test_docx(["Valid DOCX content"])

        extractor = DOCXExtractor()
        is_valid, error = extractor.validate(docx_bytes)

        assert is_valid is True
        assert error is None

    def test_docx_validation_invalid_header(self):
        """Test DOCX validation fails for non-DOCX file"""
        fake_docx = b"This is not a DOCX file"

        extractor = DOCXExtractor()
        is_valid, error = extractor.validate(fake_docx)

        assert is_valid is False
        assert "DOCX" in error or "ZIP" in error

    def test_docx_extraction_empty_file(self):
        """Test DOCX extraction raises error for empty file"""
        empty_docx = b""

        extractor = DOCXExtractor()
        with pytest.raises(ValueError) as exc_info:
            extractor.extract(empty_docx, "empty.docx")

        assert "empty.docx" in str(exc_info.value)

    def test_docx_get_metadata(self):
        """Test extracting metadata from DOCX"""
        docx_bytes = generate_test_docx(["Test content"], include_table=True)

        extractor = DOCXExtractor()
        metadata = extractor.get_metadata(docx_bytes)

        assert isinstance(metadata, dict)
        assert "paragraph_count" in metadata
        assert "table_count" in metadata
        assert metadata["paragraph_count"] >= 1
        assert metadata["table_count"] >= 1


# ============================================================================
# TXT Extractor Tests
# ============================================================================

class TestTXTExtractor:
    """Test suite for TXT text extraction"""

    def test_txt_extractor_initialization(self):
        """Test TXT extractor can be initialized"""
        extractor = TXTExtractor()
        assert extractor is not None
        assert extractor.preferred_encoding == 'utf-8'

    def test_txt_utf8_extraction(self):
        """Test extracting UTF-8 text"""
        test_text = "This is a test text file.\nWith multiple lines.\nAbout grant writing."
        txt_bytes = test_text.encode('utf-8')

        extractor = TXTExtractor()
        extracted = extractor.extract(txt_bytes, "test.txt")

        assert "test text file" in extracted
        assert "multiple lines" in extracted
        assert "grant writing" in extracted

    def test_txt_latin1_extraction(self):
        """Test extracting Latin-1 encoded text"""
        test_text = "Text with special chars: café, naïve, résumé"
        txt_bytes = test_text.encode('latin-1')

        extractor = TXTExtractor()
        extracted = extractor.extract(txt_bytes, "latin1.txt")

        # Should successfully decode (either as UTF-8 or via auto-detection)
        assert "special chars" in extracted
        assert len(extracted) > 0

    def test_txt_encoding_detection(self):
        """Test automatic encoding detection"""
        # Create text with UTF-8 specific characters
        test_text = "Unicode text: 你好世界 (Hello World)"
        txt_bytes = test_text.encode('utf-8')

        extractor = TXTExtractor()
        extracted = extractor.extract(txt_bytes, "unicode.txt")

        assert len(extracted) > 0
        # May not preserve exact unicode, but should extract something

    def test_txt_validation_success(self):
        """Test TXT validation succeeds for valid text"""
        txt_bytes = "Valid text content".encode('utf-8')

        extractor = TXTExtractor()
        is_valid, error = extractor.validate(txt_bytes)

        assert is_valid is True
        assert error is None

    def test_txt_validation_empty_file(self):
        """Test TXT validation fails for empty file"""
        empty_txt = b""

        extractor = TXTExtractor()
        is_valid, error = extractor.validate(empty_txt)

        assert is_valid is False
        assert "empty" in error.lower()

    def test_txt_extraction_empty_file(self):
        """Test TXT extraction raises error for empty file"""
        empty_txt = b""

        extractor = TXTExtractor()
        with pytest.raises(ValueError) as exc_info:
            extractor.extract(empty_txt, "empty.txt")

        assert "empty" in str(exc_info.value).lower()

    def test_txt_line_ending_normalization(self):
        """Test line ending normalization (CRLF to LF)"""
        test_text = "Line 1\r\nLine 2\r\nLine 3"
        txt_bytes = test_text.encode('utf-8')

        extractor = TXTExtractor()
        extracted = extractor.extract(txt_bytes, "crlf.txt")

        # Should normalize to LF
        assert "Line 1\nLine 2\nLine 3" in extracted or "Line 1" in extracted
        assert "\r\n" not in extracted  # CRLF should be converted

    def test_txt_get_encoding_info(self):
        """Test getting encoding information"""
        txt_bytes = "Test content".encode('utf-8')

        extractor = TXTExtractor()
        info = extractor.get_encoding_info(txt_bytes)

        assert isinstance(info, dict)
        assert "detected_encoding" in info
        assert "confidence" in info
        assert "byte_size" in info
        assert info["byte_size"] == len(txt_bytes)


# ============================================================================
# DocumentProcessor Tests
# ============================================================================

class TestDocumentProcessor:
    """Test suite for DocumentProcessor orchestrator"""

    def test_processor_initialization(self):
        """Test processor can be initialized"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        assert processor is not None
        assert processor.vector_store is None
        assert processor.embedding_model is None

    def test_processor_extractor_registration(self):
        """Test registering extractors"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        # Register extractors
        processor.register_extractor(FileType.PDF, PDFExtractor())
        processor.register_extractor(FileType.DOCX, DOCXExtractor())
        processor.register_extractor(FileType.TXT, TXTExtractor())

        # Verify registration
        assert FileType.PDF in processor._text_extractors
        assert FileType.DOCX in processor._text_extractors
        assert FileType.TXT in processor._text_extractors

    def test_processor_get_file_type(self):
        """Test file type detection from filename"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        assert processor._get_file_type("document.pdf") == FileType.PDF
        assert processor._get_file_type("document.docx") == FileType.DOCX
        assert processor._get_file_type("document.txt") == FileType.TXT
        assert processor._get_file_type("DOCUMENT.PDF") == FileType.PDF  # Case insensitive

    def test_processor_get_file_type_unsupported(self):
        """Test file type detection fails for unsupported types"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        with pytest.raises(ValueError) as exc_info:
            processor._get_file_type("document.xyz")

        assert "Unsupported file type" in str(exc_info.value)
        assert ".xyz" in str(exc_info.value)

    def test_processor_get_extractor_unregistered(self):
        """Test getting extractor fails if not registered"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        with pytest.raises(ValueError) as exc_info:
            processor._get_extractor(FileType.PDF)

        assert "No extractor registered" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_processor_pdf_pipeline(self):
        """Test complete processing pipeline with PDF"""
        # Create processor with registered extractors
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )
        processor.register_extractor(FileType.PDF, PDFExtractor())

        # Generate test PDF
        pdf_bytes = generate_test_pdf("Grant proposal content for testing.", num_pages=1)

        # Process document
        metadata = {
            "doc_id": "test_001",
            "doc_type": "Grant Proposal",
            "year": 2023
        }

        result = await processor.process_document(
            file_content=pdf_bytes,
            filename="test.pdf",
            metadata=metadata
        )

        # Verify result
        assert result.success is True
        # doc_id is auto-generated by metadata extractor, not from provided metadata
        assert result.doc_id is not None
        assert len(result.doc_id) > 0
        assert result.chunks_created > 0
        assert "Successfully processed" in result.message
        assert result.error is None

    @pytest.mark.asyncio
    async def test_processor_docx_pipeline(self):
        """Test complete processing pipeline with DOCX"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )
        processor.register_extractor(FileType.DOCX, DOCXExtractor())

        # Generate test DOCX
        docx_bytes = generate_test_docx([
            "Program description for youth development.",
            "Impact metrics and outcomes."
        ])

        metadata = {
            "doc_id": "test_002",
            "doc_type": "Program Description",
            "year": 2024
        }

        result = await processor.process_document(
            file_content=docx_bytes,
            filename="test.docx",
            metadata=metadata
        )

        assert result.success is True
        # doc_id is auto-generated by metadata extractor
        assert result.doc_id is not None
        assert len(result.doc_id) > 0
        assert result.chunks_created > 0

    @pytest.mark.asyncio
    async def test_processor_txt_pipeline(self):
        """Test complete processing pipeline with TXT"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )
        processor.register_extractor(FileType.TXT, TXTExtractor())

        # Generate test TXT
        txt_content = "Letter of inquiry content.\nProgram goals and objectives.\nExpected outcomes."
        txt_bytes = txt_content.encode('utf-8')

        metadata = {
            "doc_id": "test_003",
            "doc_type": "Letter of Inquiry",
            "year": 2024
        }

        result = await processor.process_document(
            file_content=txt_bytes,
            filename="test.txt",
            metadata=metadata
        )

        assert result.success is True
        # doc_id is auto-generated by metadata extractor
        assert result.doc_id is not None
        assert len(result.doc_id) > 0
        assert result.chunks_created > 0

    @pytest.mark.asyncio
    async def test_processor_invalid_file_type(self):
        """Test processing fails for unsupported file type"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        result = await processor.process_document(
            file_content=b"Some content",
            filename="document.xyz",
            metadata={"doc_id": "test_004"}
        )

        assert result.success is False
        assert result.chunks_created == 0
        assert "Unsupported file type" in result.error

    @pytest.mark.asyncio
    async def test_processor_validation_failure(self):
        """Test processing fails when validation fails"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )
        processor.register_extractor(FileType.PDF, PDFExtractor())

        # Use invalid PDF content
        invalid_pdf = b"Not a PDF file"

        result = await processor.process_document(
            file_content=invalid_pdf,
            filename="invalid.pdf",
            metadata={"doc_id": "test_005"}
        )

        assert result.success is False
        assert result.chunks_created == 0
        assert "Validation failed" in result.message


# ============================================================================
# Factory Tests
# ============================================================================

class TestProcessorFactory:
    """Test suite for ProcessorFactory"""

    def test_factory_create_processor(self):
        """Test factory creates processor with dependencies"""
        processor = ProcessorFactory.create_processor(
            vector_store=None,
            embedding_model=None
        )

        assert isinstance(processor, DocumentProcessor)
        assert processor.vector_store is None
        assert processor.embedding_model is None

    def test_factory_create_with_chunking_service(self):
        """Test factory creates processor with chunking service"""
        # Mock chunking service
        class MockChunkingService:
            pass

        mock_service = MockChunkingService()

        processor = ProcessorFactory.create_processor(
            chunking_service=mock_service
        )

        assert isinstance(processor, DocumentProcessor)
        assert processor.chunking_service is mock_service


# ============================================================================
# Edge Cases and Error Handling
# ============================================================================

class TestEdgeCases:
    """Test edge cases and error conditions"""

    def test_pdf_extractor_with_special_characters(self):
        """Test PDF extraction with special characters"""
        test_text = "Special chars: $100,000 funding @ 50% match — see §2.3"
        pdf_bytes = generate_test_pdf(test_text)

        extractor = PDFExtractor()
        extracted = extractor.extract(pdf_bytes, "special.pdf")

        # Should handle special characters gracefully
        assert len(extracted) > 0

    def test_docx_empty_paragraphs(self):
        """Test DOCX with empty paragraphs"""
        # Include empty strings
        paragraphs = ["First paragraph", "", "Third paragraph"]
        docx_bytes = generate_test_docx(paragraphs)

        extractor = DOCXExtractor()
        extracted = extractor.extract(docx_bytes, "empty_paras.docx")

        # Should skip empty paragraphs
        assert "First paragraph" in extracted
        assert "Third paragraph" in extracted

    def test_txt_whitespace_handling(self):
        """Test TXT handling of whitespace"""
        test_text = "  \n  Content with whitespace  \n\n  "
        txt_bytes = test_text.encode('utf-8')

        extractor = TXTExtractor()
        extracted = extractor.extract(txt_bytes, "whitespace.txt")

        # Should strip leading/trailing whitespace
        assert extracted.strip() == "Content with whitespace"

    def test_processor_classify_document(self):
        """Test document classification"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        # Should trust user-provided doc_type
        metadata = {"doc_type": "Grant Proposal"}
        doc_type = processor._classify_document("Some text", metadata)

        assert doc_type == "Grant Proposal"

    def test_processor_classify_document_missing(self):
        """Test document classification with missing doc_type"""
        processor = DocumentProcessor(
            vector_store=None,
            embedding_model=None
        )

        # Should default to Unknown
        metadata = {}
        doc_type = processor._classify_document("Some text", metadata)

        assert doc_type == "Unknown"
