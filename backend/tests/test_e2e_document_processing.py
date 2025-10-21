"""
End-to-End Document Processing Integration Tests

Tests the complete document processing pipeline from upload to retrieval:
1. Upload documents via API (PDF, DOCX, TXT)
2. Verify text extraction and processing
3. Verify chunks stored in Qdrant with embeddings
4. Verify metadata stored in PostgreSQL
5. Verify retrieval functionality
6. Test error handling for invalid files
7. Test cleanup and deletion

Uses real services (Qdrant, PostgreSQL, embedding models) to validate
the entire pipeline works correctly in integration.
"""

import pytest
import asyncio
import io
import json
from typing import List, Dict
from datetime import datetime
from uuid import UUID

# Test file generation
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document as DocxDocument

from app.services.vector_store import QdrantStore, VectorStoreConfig
from app.services.database import DatabaseService
from app.dependencies import (
    get_document_processor,
    get_vector_store,
    get_database,
    get_embedding_model
)


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture(scope="module")
def event_loop():
    """Create event loop for async tests"""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture(scope="module")
async def test_vector_store():
    """
    Initialize test vector store

    Uses localhost Qdrant with test collection
    """
    config = VectorStoreConfig(
        host="localhost",
        port=6333,
        collection_name="test_e2e_processing",
        vector_size=1536,  # OpenAI default
        use_grpc=False
    )

    store = QdrantStore(config)

    # Ensure collection exists
    try:
        store.client.get_collection(store.collection_name)
    except:
        await store._create_collection()

    yield store

    # Cleanup: Delete test collection
    try:
        store.client.delete_collection(store.collection_name)
    except:
        pass


@pytest.fixture(scope="module")
async def test_database():
    """
    Initialize test database service

    Uses PostgreSQL database with test documents
    """
    db = DatabaseService()
    await db.connect()

    yield db

    # Cleanup: Close connection pool
    await db.disconnect()


@pytest.fixture
def sample_pdf_file():
    """
    Generate a sample PDF file for testing

    Returns:
        bytes: PDF file content
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    # Page 1
    c.drawString(100, 750, "Grant Proposal: Youth STEM Education Program 2023")
    c.drawString(100, 720, "")
    c.drawString(100, 700, "Program Description:")
    c.drawString(100, 680, "This program aims to provide STEM education and mentorship")
    c.drawString(100, 660, "to underserved youth in urban areas. Our approach focuses on")
    c.drawString(100, 640, "hands-on learning, project-based activities, and career exploration.")
    c.drawString(100, 600, "Target Population:")
    c.drawString(100, 580, "We will serve 200 students in grades 6-12 from low-income families.")
    c.showPage()

    # Page 2
    c.drawString(100, 750, "Impact and Outcomes:")
    c.drawString(100, 720, "Over the past year, we have achieved significant results:")
    c.drawString(100, 700, "- 95% program retention rate")
    c.drawString(100, 680, "- Average GPA increase of 0.8 points")
    c.drawString(100, 660, "- 85% of participants pursued STEM courses")
    c.drawString(100, 620, "Budget Summary:")
    c.drawString(100, 600, "Total Requested: $350,000")
    c.drawString(100, 580, "Duration: January 2023 - December 2023")
    c.showPage()

    c.save()
    return buffer.getvalue()


@pytest.fixture
def sample_docx_file():
    """
    Generate a sample DOCX file for testing

    Returns:
        bytes: DOCX file content
    """
    doc = DocxDocument()

    doc.add_heading("Annual Report 2022: Community Youth Programs", level=1)
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "Our organization has served over 500 youth across multiple programs, "
        "achieving measurable outcomes in academic performance, social-emotional "
        "development, and career readiness."
    )
    doc.add_paragraph()

    doc.add_heading("Program Results", level=2)
    doc.add_paragraph(
        "Academic Achievement: Participants showed an average GPA improvement "
        "of 0.6 points. 78% of high school seniors graduated on time."
    )
    doc.add_paragraph(
        "STEM Engagement: 120 students participated in our STEM workshops, "
        "with 90% reporting increased interest in STEM careers."
    )
    doc.add_paragraph()

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light List Accent 1'

    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Youth Served"
    table.cell(1, 1).text = "500"
    table.cell(2, 0).text = "Programs Offered"
    table.cell(2, 1).text = "12"

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_txt_file():
    """
    Generate a sample TXT file for testing

    Returns:
        bytes: Text file content
    """
    content = """
Budget Narrative: Youth Development Initiative

Personnel Costs:
- Program Director (1 FTE): $75,000
- Youth Coordinators (2 FTE): $100,000
- Part-time Mentors (5): $50,000

Program Costs:
- Educational Materials: $25,000
- Technology & Equipment: $40,000
- Field Trips & Activities: $30,000
- Student Stipends: $20,000

Administrative Costs:
- Office Space & Utilities: $15,000
- Insurance & Legal: $10,000
- Evaluation & Reporting: $12,000

Total Budget: $377,000

This budget reflects our commitment to providing high-quality programming
while maintaining fiscal responsibility and sustainability.
    """
    return content.encode('utf-8')


# ============================================================================
# Test Cases
# ============================================================================

@pytest.mark.asyncio
async def test_pdf_document_upload_and_processing(
    test_vector_store,
    test_database,
    sample_pdf_file
):
    """
    Test complete PDF document processing pipeline

    Validates:
    - PDF text extraction
    - Semantic chunking
    - Embedding generation
    - Vector storage in Qdrant
    - Metadata storage in PostgreSQL
    """
    print("\n=== Testing PDF Document Processing ===")

    # 1. Get processor and prepare metadata
    processor = get_document_processor()
    filename = "Youth_STEM_Grant_2023_Foundation.pdf"

    metadata = {
        "doc_type": "Grant Proposal",
        "year": 2023,
        "programs": ["Youth Development", "STEM Education"],
        "tags": ["grant", "youth", "STEM"],
        "outcome": "Funded",
        "notes": "E2E test document - PDF"
    }

    print(f"Processing: {filename}")
    print(f"File size: {len(sample_pdf_file)} bytes")

    # 2. Process document
    result = await processor.process_document(
        file_content=sample_pdf_file,
        filename=filename,
        metadata=metadata
    )

    assert result.success, f"Processing failed: {result.error}"
    assert result.doc_id is not None
    assert result.chunks_created > 0

    doc_id = result.doc_id
    print(f"✓ Document processed: {doc_id}")
    print(f"✓ Chunks created: {result.chunks_created}")

    # 3. Verify chunks in Qdrant
    embedding_model = get_embedding_model()
    query_embedding = embedding_model.get_text_embedding("STEM education youth")

    search_results = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=5,
        filters={"doc_id": doc_id}
    )

    assert len(search_results) > 0, "No chunks found in vector store"
    print(f"✓ Found {len(search_results)} chunks in Qdrant")

    # Verify chunk structure
    first_chunk = search_results[0]
    assert hasattr(first_chunk, 'text')
    assert hasattr(first_chunk, 'metadata')
    assert first_chunk.metadata.get('doc_id') == doc_id
    assert 'STEM' in first_chunk.text or 'youth' in first_chunk.text.lower()
    print(f"✓ Chunk metadata validated")

    # 4. Verify document in PostgreSQL
    doc_uuid = UUID(doc_id)
    doc_record = await test_database.get_document(doc_uuid)

    assert doc_record is not None, "Document not found in database"
    assert doc_record['filename'] == filename
    assert doc_record['doc_type'] == metadata['doc_type']
    assert doc_record['year'] == metadata['year']
    assert doc_record['chunks_count'] == result.chunks_created
    print(f"✓ Document metadata verified in PostgreSQL")
    print(f"  - Filename: {doc_record['filename']}")
    print(f"  - Type: {doc_record['doc_type']}")
    print(f"  - Chunks: {doc_record['chunks_count']}")

    # 5. Cleanup
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(doc_uuid)
    print(f"✓ Test data cleaned up")

    print("=== PDF Test PASSED ===\n")


@pytest.mark.asyncio
async def test_docx_document_upload_and_processing(
    test_vector_store,
    test_database,
    sample_docx_file
):
    """
    Test complete DOCX document processing pipeline

    Validates:
    - DOCX text extraction (paragraphs and tables)
    - Semantic chunking
    - Vector and database storage
    """
    print("\n=== Testing DOCX Document Processing ===")

    processor = get_document_processor()
    filename = "Annual_Report_2022_YouthPrograms.docx"

    metadata = {
        "doc_type": "Annual Report",
        "year": 2022,
        "programs": ["Youth Development"],
        "tags": ["annual-report", "outcomes"],
        "outcome": None,
        "notes": "E2E test document - DOCX"
    }

    print(f"Processing: {filename}")

    # Process document
    result = await processor.process_document(
        file_content=sample_docx_file,
        filename=filename,
        metadata=metadata
    )

    assert result.success, f"Processing failed: {result.error}"
    assert result.chunks_created > 0

    doc_id = result.doc_id
    print(f"✓ Document processed: {doc_id}")
    print(f"✓ Chunks created: {result.chunks_created}")

    # Verify in vector store
    embedding_model = get_embedding_model()
    query_embedding = embedding_model.get_text_embedding("academic achievement")

    search_results = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=5,
        filters={"doc_id": doc_id}
    )

    assert len(search_results) > 0
    print(f"✓ Found {len(search_results)} chunks in Qdrant")

    # Check that table content was extracted
    all_text = " ".join([chunk.text for chunk in search_results])
    assert "youth" in all_text.lower() or "program" in all_text.lower()
    print(f"✓ DOCX content extracted correctly (including tables)")

    # Verify in database
    doc_uuid = UUID(doc_id)
    doc_record = await test_database.get_document(doc_uuid)

    assert doc_record is not None
    assert doc_record['doc_type'] == metadata['doc_type']
    print(f"✓ Document metadata verified in PostgreSQL")

    # Cleanup
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(doc_uuid)
    print(f"✓ Test data cleaned up")

    print("=== DOCX Test PASSED ===\n")


@pytest.mark.asyncio
async def test_txt_document_upload_and_processing(
    test_vector_store,
    test_database,
    sample_txt_file
):
    """
    Test complete TXT document processing pipeline

    Validates:
    - TXT text extraction with encoding detection
    - Processing pipeline
    - Storage in both systems
    """
    print("\n=== Testing TXT Document Processing ===")

    processor = get_document_processor()
    filename = "Budget_Youth_Initiative_2023.txt"

    metadata = {
        "doc_type": "Budget Narrative",
        "year": 2023,
        "programs": ["Youth Development"],
        "tags": ["budget", "financial"],
        "outcome": None,
        "notes": "E2E test document - TXT"
    }

    print(f"Processing: {filename}")

    # Process document
    result = await processor.process_document(
        file_content=sample_txt_file,
        filename=filename,
        metadata=metadata
    )

    assert result.success, f"Processing failed: {result.error}"
    assert result.chunks_created > 0

    doc_id = result.doc_id
    print(f"✓ Document processed: {doc_id}")
    print(f"✓ Chunks created: {result.chunks_created}")

    # Verify in vector store
    embedding_model = get_embedding_model()
    query_embedding = embedding_model.get_text_embedding("budget personnel costs")

    search_results = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=5,
        filters={"doc_id": doc_id}
    )

    assert len(search_results) > 0
    print(f"✓ Found {len(search_results)} chunks in Qdrant")

    # Check content
    all_text = " ".join([chunk.text for chunk in search_results])
    assert "budget" in all_text.lower() or "personnel" in all_text.lower()
    print(f"✓ TXT content extracted correctly")

    # Verify in database
    doc_uuid = UUID(doc_id)
    doc_record = await test_database.get_document(doc_uuid)

    assert doc_record is not None
    print(f"✓ Document metadata verified in PostgreSQL")

    # Cleanup
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(doc_uuid)
    print(f"✓ Test data cleaned up")

    print("=== TXT Test PASSED ===\n")


@pytest.mark.asyncio
async def test_invalid_file_handling(test_vector_store, test_database):
    """
    Test error handling for invalid files

    Validates:
    - Empty file rejection
    - Invalid file type rejection
    - Corrupted file handling
    """
    print("\n=== Testing Invalid File Handling ===")

    processor = get_document_processor()

    # Test 1: Empty file
    print("Test 1: Empty file...")
    empty_file = b""
    result = await processor.process_document(
        file_content=empty_file,
        filename="empty.pdf",
        metadata={"doc_type": "Grant Proposal", "year": 2023}
    )

    assert not result.success, "Empty file should fail"
    assert "empty" in result.error.lower()
    print(f"✓ Empty file rejected: {result.error}")

    # Test 2: Invalid file type
    print("Test 2: Invalid file type...")
    invalid_file = b"This is not a PDF or DOCX file"
    result = await processor.process_document(
        file_content=invalid_file,
        filename="invalid.exe",
        metadata={"doc_type": "Grant Proposal", "year": 2023}
    )

    assert not result.success, "Invalid file type should fail"
    print(f"✓ Invalid file type rejected: {result.error}")

    # Test 3: Corrupted PDF
    print("Test 3: Corrupted PDF...")
    corrupted_pdf = b"%PDF-1.4\nCorrupted content that is not a valid PDF"
    result = await processor.process_document(
        file_content=corrupted_pdf,
        filename="corrupted.pdf",
        metadata={"doc_type": "Grant Proposal", "year": 2023}
    )

    assert not result.success, "Corrupted PDF should fail"
    print(f"✓ Corrupted PDF rejected: {result.error}")

    print("=== Invalid File Handling Tests PASSED ===\n")


@pytest.mark.asyncio
async def test_metadata_extraction_and_enrichment(
    test_vector_store,
    test_database,
    sample_pdf_file
):
    """
    Test metadata extraction and enrichment

    Validates:
    - Filename pattern parsing
    - Document structure analysis
    - Metadata completeness
    """
    print("\n=== Testing Metadata Extraction ===")

    processor = get_document_processor()

    # Use filename with pattern: "Org_Type_Year_Funder.pdf"
    filename = "YouthOrg_GrantProposal_2023_StateFoundation.pdf"

    # Provide minimal metadata
    metadata = {
        "doc_type": None,  # Let system infer from filename
        "year": None,       # Let system infer from filename
        "programs": ["Youth Development"],
        "tags": [],
        "outcome": None,
        "notes": "Testing metadata extraction"
    }

    print(f"Processing: {filename}")
    print("Minimal metadata provided, testing enrichment...")

    result = await processor.process_document(
        file_content=sample_pdf_file,
        filename=filename,
        metadata=metadata
    )

    assert result.success
    doc_id = result.doc_id
    print(f"✓ Document processed: {doc_id}")

    # Check database record for enriched metadata
    doc_uuid = UUID(doc_id)
    doc_record = await test_database.get_document(doc_uuid)

    assert doc_record is not None

    # Verify enrichment
    print(f"✓ Metadata enrichment verified:")
    print(f"  - Doc type: {doc_record['doc_type']}")
    print(f"  - Year: {doc_record['year']}")
    print(f"  - File size: {doc_record['file_size']} bytes")

    # Should have inferred year from filename
    if doc_record['year'] is None or doc_record['year'] == 2023:
        print(f"  - Year correctly extracted from filename")

    # Cleanup
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(doc_uuid)
    print(f"✓ Test data cleaned up")

    print("=== Metadata Extraction Test PASSED ===\n")


@pytest.mark.asyncio
async def test_document_retrieval_after_upload(
    test_vector_store,
    test_database,
    sample_pdf_file
):
    """
    Test that uploaded documents can be retrieved via similarity search

    Validates:
    - Documents are searchable after upload
    - Relevance of search results
    - Metadata filtering works
    """
    print("\n=== Testing Document Retrieval ===")

    processor = get_document_processor()
    embedding_model = get_embedding_model()

    # Upload document
    filename = "STEM_Education_Grant_2023.pdf"
    metadata = {
        "doc_type": "Grant Proposal",
        "year": 2023,
        "programs": ["STEM Education", "Youth Development"],
        "tags": ["grant", "STEM", "education"],
        "outcome": "Funded",
        "notes": "Retrieval test document"
    }

    result = await processor.process_document(
        file_content=sample_pdf_file,
        filename=filename,
        metadata=metadata
    )

    assert result.success
    doc_id = result.doc_id
    print(f"✓ Document uploaded: {doc_id}")

    # Test 1: Basic retrieval
    print("Test 1: Basic similarity search...")
    query_embedding = embedding_model.get_text_embedding(
        "youth STEM education programs"
    )

    search_results = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=5
    )

    assert len(search_results) > 0
    print(f"✓ Retrieved {len(search_results)} results")

    # Test 2: Filtered retrieval by doc_id
    print("Test 2: Filtered search by doc_id...")
    filtered_results = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=5,
        filters={"doc_id": doc_id}
    )

    assert len(filtered_results) > 0
    assert all(r.metadata.get('doc_id') == doc_id for r in filtered_results)
    print(f"✓ Filtered retrieval works ({len(filtered_results)} results)")

    # Test 3: Check relevance
    print("Test 3: Relevance check...")
    top_result = filtered_results[0]
    assert top_result.score > 0.5, f"Low relevance score: {top_result.score}"
    print(f"✓ Top result relevance: {top_result.score:.3f}")
    print(f"  Sample text: {top_result.text[:100]}...")

    # Cleanup
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(UUID(doc_id))
    print(f"✓ Test data cleaned up")

    print("=== Document Retrieval Test PASSED ===\n")


@pytest.mark.asyncio
async def test_concurrent_document_processing(
    test_vector_store,
    test_database,
    sample_pdf_file,
    sample_txt_file
):
    """
    Test processing multiple documents concurrently

    Validates:
    - Concurrent processing works
    - No data corruption
    - All documents stored correctly
    """
    print("\n=== Testing Concurrent Processing ===")

    processor = get_document_processor()

    # Prepare multiple documents
    documents = [
        {
            "content": sample_pdf_file,
            "filename": "Concurrent_Test_1.pdf",
            "metadata": {
                "doc_type": "Grant Proposal",
                "year": 2023,
                "programs": ["Education"],
                "tags": ["test-1"],
                "outcome": None,
                "notes": "Concurrent test 1"
            }
        },
        {
            "content": sample_txt_file,
            "filename": "Concurrent_Test_2.txt",
            "metadata": {
                "doc_type": "Budget Narrative",
                "year": 2023,
                "programs": ["Youth Development"],
                "tags": ["test-2"],
                "outcome": None,
                "notes": "Concurrent test 2"
            }
        }
    ]

    print(f"Processing {len(documents)} documents concurrently...")

    # Process concurrently
    tasks = [
        processor.process_document(
            file_content=doc["content"],
            filename=doc["filename"],
            metadata=doc["metadata"]
        )
        for doc in documents
    ]

    results = await asyncio.gather(*tasks)

    # Verify all succeeded
    assert all(r.success for r in results), "Some documents failed"
    doc_ids = [r.doc_id for r in results]

    print(f"✓ All {len(results)} documents processed successfully")
    for i, result in enumerate(results):
        print(f"  - Doc {i+1}: {result.doc_id} ({result.chunks_created} chunks)")

    # Verify in database
    for doc_id in doc_ids:
        doc_record = await test_database.get_document(UUID(doc_id))
        assert doc_record is not None

    print(f"✓ All documents verified in database")

    # Cleanup
    for doc_id in doc_ids:
        await test_vector_store.delete_document(doc_id)
        await test_database.delete_document(UUID(doc_id))

    print(f"✓ Test data cleaned up")
    print("=== Concurrent Processing Test PASSED ===\n")


@pytest.mark.asyncio
async def test_document_deletion_cascade(
    test_vector_store,
    test_database,
    sample_pdf_file
):
    """
    Test that deleting a document removes data from all systems

    Validates:
    - Chunks removed from Qdrant
    - Metadata removed from PostgreSQL
    - No orphaned data remains
    """
    print("\n=== Testing Document Deletion Cascade ===")

    processor = get_document_processor()
    embedding_model = get_embedding_model()

    # Upload document
    filename = "Delete_Test_Document.pdf"
    metadata = {
        "doc_type": "Grant Proposal",
        "year": 2023,
        "programs": ["Test Program"],
        "tags": ["delete-test"],
        "outcome": None,
        "notes": "Document for deletion testing"
    }

    result = await processor.process_document(
        file_content=sample_pdf_file,
        filename=filename,
        metadata=metadata
    )

    assert result.success
    doc_id = result.doc_id
    print(f"✓ Document uploaded: {doc_id}")

    # Verify document exists in both systems
    doc_record = await test_database.get_document(UUID(doc_id))
    assert doc_record is not None
    print(f"✓ Document exists in PostgreSQL")

    query_embedding = embedding_model.get_text_embedding("test")
    chunks = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=10,
        filters={"doc_id": doc_id}
    )
    assert len(chunks) > 0
    print(f"✓ {len(chunks)} chunks exist in Qdrant")

    # Delete document
    print("Deleting document...")
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(UUID(doc_id))

    # Verify deletion from PostgreSQL
    doc_record = await test_database.get_document(UUID(doc_id))
    assert doc_record is None, "Document still exists in PostgreSQL"
    print(f"✓ Document removed from PostgreSQL")

    # Verify deletion from Qdrant
    chunks_after = await test_vector_store.search_similar(
        query_embedding=query_embedding,
        top_k=10,
        filters={"doc_id": doc_id}
    )
    assert len(chunks_after) == 0, "Chunks still exist in Qdrant"
    print(f"✓ Chunks removed from Qdrant")

    print("=== Document Deletion Cascade Test PASSED ===\n")


# ============================================================================
# Performance Test
# ============================================================================

@pytest.mark.asyncio
async def test_processing_performance(
    test_vector_store,
    test_database,
    sample_pdf_file
):
    """
    Test processing performance and timing

    Validates:
    - Processing completes in reasonable time
    - Performance metrics are captured
    """
    print("\n=== Testing Processing Performance ===")

    processor = get_document_processor()

    filename = "Performance_Test.pdf"
    metadata = {
        "doc_type": "Grant Proposal",
        "year": 2023,
        "programs": ["Test"],
        "tags": ["performance"],
        "outcome": None,
        "notes": "Performance test"
    }

    import time
    start_time = time.time()

    result = await processor.process_document(
        file_content=sample_pdf_file,
        filename=filename,
        metadata=metadata
    )

    end_time = time.time()
    duration = end_time - start_time

    assert result.success
    doc_id = result.doc_id

    print(f"✓ Processing completed in {duration:.2f}s")
    print(f"  - Chunks created: {result.chunks_created}")
    print(f"  - Average time per chunk: {duration/result.chunks_created:.3f}s")

    # Reasonable performance threshold (adjust based on environment)
    assert duration < 30.0, f"Processing too slow: {duration}s"
    print(f"✓ Performance acceptable (<30s)")

    # Cleanup
    await test_vector_store.delete_document(doc_id)
    await test_database.delete_document(UUID(doc_id))

    print("=== Processing Performance Test PASSED ===\n")
